"""
文件解析模块 - 统一处理用户上传的各种文件格式

支持以下文件类型：
- 图片（jpg/png/jpeg/gif/webp）→ 用智谱GLM-4V识别
- PDF（.pdf）                    → 用PyPDF2提取文字
- Excel（.xlsx/.xls）            → 用openpyxl提取表格文字
- Word（.docx）                  → 用python-docx提取文字
- 文本（.txt/.md/.csv）          → 直接读取

核心方法：parse_uploaded_file(file_bytes, filename)
返回：(file_type, content) - 文件类型和内容（图片返回base64，其他返回文字）
"""

import os
import base64
import io


# 支持的文件类型常量
TYPE_IMAGE = "image"      # 图片类型（用智谱GLM-4V识别）
TYPE_TEXT = "text"        # 文本类型（直接是文字内容）


def parse_uploaded_file(file_bytes, filename):
    """
    解析上传的文件 - 根据文件扩展名自动选择解析方式

    参数：
        file_bytes: 文件的二进制数据
        filename: 文件名（用于判断文件类型）

    返回：字典 {"type": "image"或"text", "content": 内容}
        - 图片：type="image", content=base64字符串
        - 文本类：type="text", content=提取的文字
    """
    # 获取文件扩展名（转小写方便比较）
    ext = os.path.splitext(filename)[1].lower()

    # ---------- 图片类型 ----------
    if ext in [".jpg", ".jpeg", ".png", ".gif", ".webp", ".bmp"]:
        image_base64 = base64.b64encode(file_bytes).decode("utf-8")
        return {"type": TYPE_IMAGE, "content": image_base64}

    # ---------- PDF文档 ----------
    if ext == ".pdf":
        text = _extract_pdf_text(file_bytes)
        return {"type": TYPE_TEXT, "content": text}

    # ---------- Excel文档 ----------
    if ext in [".xlsx", ".xls"]:
        text = _extract_excel_text(file_bytes)
        return {"type": TYPE_TEXT, "content": text}

    # ---------- Word文档 ----------
    if ext == ".docx":
        text = _extract_docx_text(file_bytes)
        return {"type": TYPE_TEXT, "content": text}

    # ---------- 纯文本类 ----------
    if ext in [".txt", ".md", ".csv", ".json"]:
        text = file_bytes.decode("utf-8", errors="ignore")
        return {"type": TYPE_TEXT, "content": text}

    # 不支持的类型，尝试当文本读
    try:
        text = file_bytes.decode("utf-8", errors="ignore")
        return {"type": TYPE_TEXT, "content": text}
    except Exception:
        raise Exception(f"不支持的文件格式：{ext}")


def _extract_pdf_text(file_bytes):
    """从PDF文件中提取文字内容"""
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(io.BytesIO(file_bytes))
        text = ""
        # 逐页提取文字
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
        return text.strip() if text else "（PDF未提取到文字，可能是扫描版PDF，建议转成图片上传）"
    except Exception as e:
        return f"（PDF解析失败：{str(e)}）"


def _extract_excel_text(file_bytes):
    """从Excel文件中提取表格文字内容"""
    try:
        from openpyxl import load_workbook
        wb = load_workbook(io.BytesIO(file_bytes), data_only=True)
        text = ""
        # 遍历所有工作表
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            text += f"【工作表：{sheet_name}】\n"
            # 遍历所有行，把每个单元格的值拼成文本
            for row in ws.iter_rows(values_only=True):
                # 过滤掉全空的行
                if any(cell is not None for cell in row):
                    row_text = " | ".join(str(cell) if cell is not None else "" for cell in row)
                    text += row_text + "\n"
            text += "\n"
        return text.strip()
    except Exception as e:
        return f"（Excel解析失败：{str(e)}）"


def _extract_docx_text(file_bytes):
    """从Word文档中提取文字内容"""
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        text = ""
        # 提取所有段落文字
        for para in doc.paragraphs:
            if para.text.strip():
                text += para.text + "\n"
        # 提取所有表格内容
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text for cell in row.cells)
                text += row_text + "\n"
        return text.strip() if text else "（Word文档未提取到文字）"
    except Exception as e:
        return f"（Word解析失败：{str(e)}）"
